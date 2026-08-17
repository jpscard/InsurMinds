# utils/processing.py

import pandas as pd
import zipfile
import io
import re
import unicodedata
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import streamlit as st


def limpar_nomes_colunas(df: pd.DataFrame) -> pd.DataFrame:
    """Padroniza os nomes das colunas removendo acentos, espaços e caracteres especiais."""
    cols_novas = []
    for col in df.columns:
        col_str = str(col)
        col_norm = ''.join(c for c in unicodedata.normalize('NFD', col_str) if unicodedata.category(c) != 'Mn')
        col_norm = col_norm.lower().strip().replace(' ', '_').replace('/', '_').replace('-', '_')
        col_norm = col_norm.replace('(', '').replace(')', '').replace('.', '').replace('?', '')
        cols_novas.append(col_norm)
    df.columns = cols_novas
    return df


def detectar_formato_csv(stream_bytes: bytes) -> tuple[str, str, str]:
    """
    Detecta automaticamente o encoding, o delimitador (sep) e o caractere decimal do CSV.
    Retorna: (encoding, sep, decimal)
    """
    # 1. Detectar encoding
    encodings_to_try = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
    sample_text = ""
    detected_encoding = 'utf-8'
    for enc in encodings_to_try:
        try:
            sample_text = stream_bytes[:8192].decode(enc)
            detected_encoding = enc
            break
        except (UnicodeDecodeError, LookupError):
            continue

    if not sample_text:
        sample_text = stream_bytes[:8192].decode('latin-1', errors='ignore')
        detected_encoding = 'latin-1'

    # 2. Detectar separador
    lines = [l for l in sample_text.splitlines() if l.strip()][:5]
    first_line = lines[0] if lines else sample_text
    
    count_semicolon = first_line.count(';')
    count_comma = first_line.count(',')
    count_tab = first_line.count('\t')
    count_pipe = first_line.count('|')

    counts = {';': count_semicolon, ',': count_comma, '\t': count_tab, '|': count_pipe}
    sep = max(counts, key=counts.get)
    if counts[sep] == 0:
        sep = ','

    # 3. Detectar decimal
    # Se o separador for ';', números fracionários no Brasil costumam usar ',' (ex: "4603,42")
    # Se o separador for ',', números costumam usar '.'
    decimal = '.'
    if sep == ';':
        if re.search(r'\d+,\d{2}', sample_text):
            decimal = ','
    else:
        if re.search(r'\d+\.\d{2}', sample_text):
            decimal = '.'
        elif re.search(r'\d+,\d{2}', sample_text):
            decimal = ','

    return detected_encoding, sep, decimal


def ler_csv_bytes(bytes_content: bytes, nome_arquivo: str = "") -> pd.DataFrame:
    """Lê bytes de um arquivo CSV com detecção automática de formato e limpeza."""
    encoding, sep, decimal = detectar_formato_csv(bytes_content)
    
    try:
        df = pd.read_csv(
            io.BytesIO(bytes_content),
            sep=sep,
            decimal=decimal,
            encoding=encoding,
            low_memory=False
        )
    except Exception:
        # Fallback de leitura permissiva
        df = pd.read_csv(
            io.BytesIO(bytes_content),
            sep=sep,
            encoding='latin-1',
            on_bad_lines='skip',
            low_memory=False
        )

    df = limpar_nomes_colunas(df)
    
    colunas_preservar_str = ['chave', 'cpf', 'cnpj', 'codigo', 'protocolo', 'recibo', 'nfe', 'nfs', 'id', 'barras', 'telefone', 'cep']
    for col in df.columns:
        if any(p in col for p in colunas_preservar_str):
            df[col] = df[col].astype(str).str.replace(r'\.0$', '', regex=True)
            continue
        if df[col].dtype == object:
            # Tenta converter strings numéricas no padrão pt-BR ("1.234,56" ou "1234,56") se aplicável
            sample = df[col].dropna().astype(str).head(20)
            if sample.str.match(r'^-?\d+([.,]\d+)?$').all() and len(sample) > 0:
                try:
                    cleaned_series = df[col].astype(str).str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
                    df[col] = pd.to_numeric(cleaned_series, errors='ignore')
                except Exception:
                    pass

    return df


def ler_csv_flexivel(zip_ref: zipfile.ZipFile, nome_arquivo: str) -> pd.DataFrame:
    """Lê um arquivo CSV de dentro do ZIP com detecção automática de formato e limpeza."""
    with zip_ref.open(nome_arquivo) as f:
        bytes_content = f.read()
    return ler_csv_bytes(bytes_content, nome_arquivo)


def extrair_dicionario_dados(zip_ref: zipfile.ZipFile) -> str | None:
    """
    Localiza e extrai qualquer arquivo de documentação/dicionário de dados dentro do ZIP.
    Ex: dicionario.txt, dicionario.csv, schema.json, readme.md, etc.
    """
    nomes = zip_ref.namelist()
    palavras_chave_doc = ['dicionario', 'dictionary', 'schema', 'descricao', 'metadados', 'metadata', 'readme', 'leia_me']
    
    candidatos_doc = []
    for nome in nomes:
        if '__MACOSX' in nome or nome.endswith('/'):
            continue
        nome_lower = nome.lower()
        if any(kw in nome_lower for kw in palavras_chave_doc) or nome_lower.endswith(('.txt', '.md', '.json', '.yaml', '.yml')):
            score = 2 if any(kw in nome_lower for kw in palavras_chave_doc) else 1
            candidatos_doc.append((score, nome))

    if not candidatos_doc:
        return None

    candidatos_doc.sort(key=lambda x: x[0], reverse=True)
    melhor_arquivo = candidatos_doc[0][1]

    try:
        with zip_ref.open(melhor_arquivo) as f:
            conteudo_bytes = f.read()
        
        for enc in ['utf-8', 'latin-1', 'cp1252']:
            try:
                texto = conteudo_bytes.decode(enc)
                return f"=== Dicionário de Dados ({melhor_arquivo}) ===\n{texto.strip()}"
            except UnicodeDecodeError:
                continue
    except Exception as e:
        print(f"Erro ao extrair dicionário de dados: {e}")
        return None

    return None


def processar_arquivos(arquivos_input) -> tuple[pd.DataFrame, str | None, list[str]]:
    """
    Processa arquivos compactados (.ZIP) ou múltiplos arquivos (.CSV, .TXT, etc.):
    - Suporta arquivos .ZIP com múltiplos CSVs internos.
    - Suporta múltiplos arquivos CSV e Dicionários enviados simultaneamente.
    - Realiza merge automático inteligente entre tabelas de Cabeçalho e Itens.
    - Extrai dicionário de dados se presente no ZIP ou nos uploads.
    - Converte colunas de data e números automaticamente.
    
    Retorna: (df_completo, texto_dicionario, lista_arquivos_processados)
    """
    try:
        if not isinstance(arquivos_input, (list, tuple)):
            arquivos_input = [arquivos_input]

        dfs = {}
        dicionarios = []
        nomes_arquivos_processados = []

        for item in arquivos_input:
            if hasattr(item, 'name'):
                nome_item = item.name
                bytes_item = item.read() if hasattr(item, 'read') else open(item, 'rb').read()
            elif isinstance(item, str):
                nome_item = item
                with open(item, 'rb') as f:
                    bytes_item = f.read()
            else:
                nome_item = getattr(item, 'name', 'arquivo.zip')
                bytes_item = item.read()

            nome_item_lower = nome_item.lower()

            if nome_item_lower.endswith('.zip') or zipfile.is_zipfile(io.BytesIO(bytes_item)):
                with zipfile.ZipFile(io.BytesIO(bytes_item), 'r') as z:
                    dict_txt = extrair_dicionario_dados(z)
                    if dict_txt:
                        dicionarios.append(dict_txt)

                    for n in z.namelist():
                        if n.startswith('__MACOSX') or n.endswith('/'):
                            continue
                        if n.lower().endswith('.csv'):
                            with z.open(n) as zf:
                                df_csv = ler_csv_bytes(zf.read(), n)
                                dfs[n] = df_csv
                                nomes_arquivos_processados.append(n)

            elif nome_item_lower.endswith('.csv'):
                df_csv = ler_csv_bytes(bytes_item, nome_item)
                dfs[nome_item] = df_csv
                nomes_arquivos_processados.append(nome_item)

            elif nome_item_lower.endswith(('.txt', '.md', '.json', '.yaml', '.yml')):
                for enc in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        txt = bytes_item.decode(enc)
                        dicionarios.append(f"=== Dicionário de Dados ({nome_item}) ===\n{txt.strip()}")
                        break
                    except UnicodeDecodeError:
                        continue

        if not dfs:
            raise ValueError("Nenhum arquivo CSV válido foi encontrado nos arquivos fornecidos.")

        dicionario_texto = "\n\n".join(dicionarios) if dicionarios else None
        arquivos_csv = list(dfs.keys())

        # Se houver apenas 1 CSV
        if len(arquivos_csv) == 1:
            df = dfs[arquivos_csv[0]]
            for col in df.columns:
                if 'data' in col and not col.endswith('_timestamp'):
                    df[col] = pd.to_datetime(df[col], errors='coerce')
            return df, dicionario_texto, nomes_arquivos_processados

        # Se houver múltiplos CSVs (ex: Cabeçalho/NotaFiscal e Itens/NotaFiscalItem)
        cabecalho_nome = None
        itens_nome = None

        for nome in arquivos_csv:
            nome_lower = nome.lower()
            if any(k in nome_lower for k in ['cabecalho', 'notafiscal.', 'nota_fiscal.', 'master', 'header']) and not any(k in nome_lower for k in ['item', 'itens', 'item.', 'detail']):
                cabecalho_nome = nome
            elif any(k in nome_lower for k in ['item', 'itens', 'notafiscalitem', 'detail']):
                itens_nome = nome

        # Caso não tenha identificado pelo nome, usar a quantidade de linhas (itens > cabecalho)
        if not cabecalho_nome or not itens_nome or cabecalho_nome == itens_nome:
            ordenados_por_tamanho = sorted(arquivos_csv, key=lambda k: len(dfs[k]))
            cabecalho_nome = ordenados_por_tamanho[0]
            itens_nome = ordenados_por_tamanho[-1]

        df_cabecalho = dfs[cabecalho_nome]
        df_itens = dfs[itens_nome]

        # Identificar chave de junção comum
        colunas_comuns = [c for c in df_cabecalho.columns if c in df_itens.columns]
        chave_merge = None
        for candidata in ['chave_de_acesso', 'chave', 'id', 'numero_nota', 'codigo', 'numero']:
            if candidata in colunas_comuns:
                chave_merge = candidata
                break
        
        if not chave_merge and colunas_comuns:
            chave_merge = colunas_comuns[0]

        if chave_merge:
            # Remove colunas duplicadas do cabeçalho que já estão em itens para evitar _x e _y excessivos
            cols_cabecalho = [c for c in df_cabecalho.columns if c not in df_itens.columns or c == chave_merge]
            df_completo = pd.merge(df_cabecalho[cols_cabecalho], df_itens, on=chave_merge, how='inner')
        else:
            # Se não houver chave comum identificável, concatena colunas
            df_completo = pd.concat([df_cabecalho, df_itens], axis=1)

        # Normalização de nomes de colunas úteis para compatibilidade ampla
        for col in list(df_completo.columns):
            if col.endswith('_x'):
                base_name = col[:-2]
                if base_name not in df_completo.columns:
                    df_completo[base_name] = df_completo[col]
            elif not col.endswith('_x') and not col.endswith('_y'):
                alias_x = f"{col}_x"
                if alias_x not in df_completo.columns:
                    df_completo[alias_x] = df_completo[col]

        # Converte datas
        for col in df_completo.columns:
            if 'data' in col and not col.endswith('_timestamp'):
                df_completo[col] = pd.to_datetime(df_completo[col], errors='coerce')

        return df_completo, dicionario_texto, nomes_arquivos_processados

    except Exception as e:
        raise RuntimeError(f"Falha ao processar os arquivos: {e}") from e


def processar_zip(arquivo_zip) -> tuple[pd.DataFrame, str | None, list[str]]:
    """Função legada mantida para compatibilidade retroativa."""
    return processar_arquivos(arquivo_zip)


@st.cache_data
def criar_documento_word(report_items):
    """
    Gera um documento Word profissional e bem formatado com os itens selecionados.
    """
    document = Document()
    
    styles = document.styles
    styles['Title'].font.name = 'Calibri'
    styles['Title'].font.size = Pt(26)
    styles['Heading 1'].font.name = 'Calibri'
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 2'].font.name = 'Calibri'
    styles['Heading 2'].font.size = Pt(13)

    # Cabeçalho do Relatório
    document.add_heading('Relatório de Análise Inteligente de Dados', level=0)
    p_data = document.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_geracao = datetime.now().strftime("%d de %B de %Y, %H:%M:%S")
    p_data.add_run(f'Relatório gerado em: {data_geracao}').italic = True
    document.add_paragraph()

    # Renderização dos Itens do Relatório
    for item in report_items:
        try:
            titulo_item = item.get('title', 'Item de Relatório')
            document.add_heading(titulo_item, level=2)
            content = item['content']
            
            if item['type'] == 'qa' or item.get('category') == 'insight_ia':
                document.add_paragraph(f"Pergunta: {content['pergunta']}", style='Intense Quote')
                document.add_paragraph(f"Resposta: {content['resposta']}")
                if content.get('imagens'):
                    for img_bytes in content['imagens']:
                        img_buf = io.BytesIO(img_bytes)
                        document.add_picture(img_buf, width=Inches(6.0))
                        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif item['type'] == 'summary':
                document.add_paragraph(content.get('texto', ''))

            elif item['type'] == 'dataframe':
                df_item = content['dados']
                if df_item.index.name is not None:
                    df_item = df_item.reset_index()
                if not df_item.empty:
                    t = document.add_table(df_item.shape[0] + 1, df_item.shape[1], style='Table Grid')
                    for j, col_name in enumerate(df_item.columns):
                        cell = t.cell(0, j)
                        cell.text = str(col_name)
                        cell.paragraphs[0].runs[0].font.bold = True
                    for i in range(df_item.shape[0]):
                        for j in range(df_item.shape[1]):
                            valor = df_item.values[i, j]
                            texto = f"{valor:,.2f}" if isinstance(valor, (int, float)) else str(valor)
                            t.cell(i + 1, j).text = texto
                else:
                    document.add_paragraph("Nenhum dado para exibir nesta análise.")

            elif item['type'] == 'chart':
                fig = content['fig']
                fig.update_layout(template='plotly_white')
                img_buffer = io.BytesIO()
                fig.write_image(img_buffer, format='png', width=900, height=500, scale=2)
                img_buffer.seek(0)
                document.add_picture(img_buffer, width=Inches(6.5))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph()

        except Exception as e:
            print(f"ERRO AO PROCESSAR ITEM PARA DOCX: {item.get('title', 'N/A')}. Detalhes: {e}")
            document.add_paragraph(f"Não foi possível renderizar o item: {item.get('title', 'N/A')}", style='Body Text')

    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer